import PyPDF2
import docx
import os

def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                text += reader.pages[page_num].extract_text() or ""
    except Exception as e:
        print(f"Error extracting text from PDF {pdf_path}: {e}")
    return text

def extract_text_from_docx(docx_path):
    text = ""
    try:
        document = docx.Document(docx_path)
        for para in document.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"Error extracting text from DOCX {docx_path}: {e}")
    return text

def extract_text(file_path):
    if not os.path.exists(file_path):
        print(f"Error: File not found at {file_path}")
        return ""

    if file_path.lower().endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.lower().endswith('.docx'):
        return extract_text_from_docx(file_path)
    elif file_path.lower().endswith('.txt'): # Add this block for .txt files
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            print(f"Error extracting text from TXT {file_path}: {e}")
            return ""
    else:
        print(f"Unsupported file type: {os.path.basename(file_path)}. Only .pdf, .docx, and .txt are supported.")
        return ""

if __name__ == "__main__":
    # Create dummy files for testing
    # For PDF, you'll need a sample PDF file. You can create a simple one or use an existing one.
    # For DOCX, you can create a simple Word document.

    # Example: Create a dummy DOCX file
    try:
        doc = docx.Document()
        doc.add_paragraph("This is a dummy DOCX file for testing.")
        doc.add_paragraph("It contains some sample text.")
        doc.save("dummy.docx")
        print("Created dummy.docx")
    except Exception as e:
        print(f"Could not create dummy.docx: {e}")

    # Example: For PDF, you'll need to manually place a dummy.pdf in the directory
    # For demonstration, let's assume dummy.pdf exists.
    # If you don't have one, you can skip this part or create one using other tools.

    print("\n--- Testing Text Extraction ---")

    # Test with dummy.pdf (if it exists)
    pdf_test_path = "dummy.pdf"
    if os.path.exists(pdf_test_path):
        print(f"\nExtracting from {pdf_test_path}:")
        pdf_text = extract_text(pdf_test_path)
        print(pdf_text)
    else:
        print(f"\nSkipping PDF test: {pdf_test_path} not found. Please create a dummy.pdf for testing.")

    # Test with dummy.docx
    docx_test_path = "dummy.docx"
    print(f"\nExtracting from {docx_test_path}:")
    docx_text = extract_text(docx_test_path)
    print(docx_text)

    # Test with an unsupported file type
    unsupported_file_path = "dummy.txt"
    with open(unsupported_file_path, "w") as f:
        f.write("This is a dummy text file.")
    print(f"\nExtracting from {unsupported_file_path}:")
    unsupported_text = extract_text(unsupported_file_path)
    print(unsupported_text)
    os.remove(unsupported_file_path)

    # Test with a non-existent file
    non_existent_file_path = "non_existent.pdf"
    print(f"\nExtracting from {non_existent_file_path}:")
    non_existent_text = extract_text(non_existent_file_path)
    print(non_existent_text)